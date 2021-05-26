from typing import Text
from docx import Document
from docx.shared import Inches
from docx.enum.style import WD_STYLE_TYPE
import pyttsx3

#speak declaration
def speak(text):
    pyttsx3.speak(text)

#file cv.docx created
document = Document()

#profile picture
document.add_picture(
    'blind_beetle.jpg',
    width = Inches(1.5)
)

#personal information
name = input('What is your name?')
speak('Hello,' + name + '. And welcome to our C.V. generator system.')
phone_number = input('What is your phone number?')
email = input('What is your email?')

document.add_paragraph(
    name + ' | ' + phone_number + ' | ' + email
)

#simple data about yourself
document.add_heading('About me')
document.add_paragraph( input('tell me about yourself: '))

#Work Experiences
document.add_heading('Work Experience')
p = document.add_paragraph()

company = input('Enter company: ')
from_date = input('from date: ')
to_date = input('to date: ')

p.add_run(company + ' ' ).bold = True
p.add_run(from_date + ' - ' + to_date + '\n').italic = True
experience_details = input(
    'Describe your experience at ' + company + ': ')
p.add_run(experience_details)

#More experiences
while True:
    has_more_experences = input('Do you have more experiences? \nyes (type y) or no (type n): ')
    if has_more_experences.lower() == 'y':
        p = document.add_paragraph()

        company = input('Enter company: ')
        from_date = input('from date: ')
        to_date = input('to date: ')

        p.add_run(company + ' ' ).bold = True
        p.add_run(from_date + ' - ' + to_date + '\n').italic = True
        experience_details = input(
            'Describe your experience at ' + company + ': ')
        p.add_run(experience_details)
    
    elif has_more_experences.lower() == 'n':
        break

#Skills
document.add_heading('Skills') #simple data about yourself

add_skill = input(
    'Would you like to add a skill? \n type the skill or "n" if you dont want to add more skills')

while True:

    p = document.add_paragraph()#every bullet has to be a new paragraph
    
    if add_skill.lower() == 'n' :
        break
    else:
        p.style = "List Bullet"
        p.add_run(add_skill)
        add_skill = input('type skill or "n": ')

#footer
section = document.section[0]
footer = section.footer
p = footer.paragraphs[0]
p.text = "CV generated using CV-Moly-Gen"

document.save('cv.docx') #file cv.docx saved
